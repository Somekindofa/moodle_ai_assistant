"""Course RAG service — per-course ChromaDB collections with semantic chunking."""

import base64
import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

from langchain_chroma import Chroma
from langchain_core.documents.base import Document
from langchain_openai import OpenAIEmbeddings

from config.settings import ConfigurationManager
from services import translation_service

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────
# Semantic chunker
# ─────────────────────────────────────────────────────────────────

HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
TARGET_TOKENS = 400   # ~1600 chars
MIN_TOKENS = 50       # ~200 chars — merge if below this


def _approx_tokens(text: str) -> int:
    """Rough token estimate: 4 chars ≈ 1 token."""
    return max(1, len(text) // 4)


class SemanticChunker:
    """Splits structured documents (HTML / PDF / DOCX) into semantic chunks.

    Each chunk captures a coherent section bounded by headings and paragraph
    breaks.  The heading breadcrumb is prepended to the chunk text so that
    vector search is hierarchy-aware.
    """

    # ── HTML ──────────────────────────────────────────────────────

    def chunk_html(
        self,
        html: str,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:
        """Parse HTML and return semantic Document chunks."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("beautifulsoup4 not installed — cannot chunk HTML")
            return []

        soup = BeautifulSoup(html, "html.parser")
        return self._walk_soup(soup, base_metadata)

    def _walk_soup(
        self, soup: Any, base_metadata: Dict[str, Any]
    ) -> List[Document]:
        from bs4 import NavigableString, Tag

        heading_stack: List[str] = []   # breadcrumb of current headings
        buffer: List[str] = []           # accumulated paragraph text
        chunks: List[Document] = []

        def flush(force: bool = False) -> None:
            text = " ".join(buffer).strip()
            if not text:
                buffer.clear()
                return
            if not force and _approx_tokens(text) < MIN_TOKENS:
                return  # too small — keep accumulating
            _emit(text)
            buffer.clear()

        def _emit(text: str) -> None:
            heading_path = " > ".join(heading_stack) if heading_stack else ""
            # Prepend breadcrumb so the embedding encodes positional context
            full_text = f"{heading_path}\n\n{text}".strip() if heading_path else text
            idx = len(chunks)
            meta = {
                **base_metadata,
                "chunk_index": idx,
                "heading_path": heading_path,
                "source": f"course_{base_metadata.get('course_id')}_module_{base_metadata.get('module_id')}_chunk_{idx}",
            }
            chunks.append(Document(page_content=full_text, metadata=meta))

        for element in soup.descendants:
            if not isinstance(element, (NavigableString,)) and not isinstance(element, type(soup)):
                tag_name = getattr(element, "name", None)
                if tag_name is None:
                    continue

                if tag_name in HEADING_TAGS:
                    flush(force=True)
                    level = int(tag_name[1])
                    heading_text = element.get_text(" ", strip=True)
                    # Keep heading_stack up to (level - 1) then add new heading
                    heading_stack[:] = heading_stack[: level - 1]
                    heading_stack.append(heading_text)

                elif tag_name in ("p", "li", "td", "th", "dd", "dt", "blockquote"):
                    text = element.get_text(" ", strip=True)
                    if text:
                        buffer.append(text)
                        # Flush if buffer is large enough
                        combined = " ".join(buffer)
                        if _approx_tokens(combined) >= TARGET_TOKENS:
                            flush(force=True)

        flush(force=True)
        return chunks

    # ── PDF ───────────────────────────────────────────────────────

    def chunk_pdf(
        self,
        file_bytes: bytes,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:
        """Extract text from PDF and return semantic Document chunks."""
        try:
            import fitz  # pymupdf
        except ImportError:
            logger.error("pymupdf not installed — cannot chunk PDF")
            return []

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as e:
            logger.error(f"Failed to open PDF: {e}")
            return []

        heading_stack: List[str] = []
        buffer: List[str] = []
        chunks: List[Document] = []
        page_images: Dict[int, int] = {}   # page_num → image count

        def flush(force: bool = False) -> None:
            text = " ".join(buffer).strip()
            if not text:
                buffer.clear()
                return
            if not force and _approx_tokens(text) < MIN_TOKENS:
                return
            heading_path = " > ".join(heading_stack) if heading_stack else ""
            full_text = f"{heading_path}\n\n{text}".strip() if heading_path else text
            idx = len(chunks)
            meta = {
                **base_metadata,
                "chunk_index": idx,
                "heading_path": heading_path,
                "source": f"course_{base_metadata.get('course_id')}_module_{base_metadata.get('module_id')}_chunk_{idx}",
            }
            if page_images:
                meta["image_pages"] = ",".join(str(p) for p in sorted(page_images))
            chunks.append(Document(page_content=full_text, metadata=meta))
            buffer.clear()
            page_images.clear()

        # Heuristic: the median font size of body text
        all_sizes: List[float] = []
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") == 0:  # text block
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            all_sizes.append(span.get("size", 12))

        if all_sizes:
            all_sizes.sort()
            body_size = all_sizes[len(all_sizes) // 2]  # median
        else:
            body_size = 12.0

        heading_threshold = body_size * 1.15  # 15% larger than body = heading

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                btype = block.get("type")
                if btype == 1:  # image block
                    page_images[page_num] = page_images.get(page_num, 0) + 1
                    continue
                if btype != 0:
                    continue

                for line in block.get("lines", []):
                    line_text = " ".join(
                        span["text"] for span in line.get("spans", [])
                    ).strip()
                    if not line_text:
                        continue

                    # Check if this line looks like a heading
                    span_sizes = [s.get("size", 12) for s in line.get("spans", [])]
                    avg_size = sum(span_sizes) / len(span_sizes) if span_sizes else 12
                    is_bold = any(
                        s.get("flags", 0) & 0b10000 for s in line.get("spans", [])
                    )

                    if avg_size >= heading_threshold or (is_bold and len(line_text) < 80):
                        flush(force=True)
                        # Approximate heading level from font size
                        level = max(1, min(3, round((heading_threshold - avg_size + heading_threshold) / heading_threshold)))
                        heading_stack[:] = heading_stack[:level - 1]
                        heading_stack.append(line_text)
                    else:
                        buffer.append(line_text)
                        combined = " ".join(buffer)
                        if _approx_tokens(combined) >= TARGET_TOKENS:
                            flush(force=True)

        flush(force=True)
        return chunks

    # ── DOCX ──────────────────────────────────────────────────────

    def chunk_docx(
        self,
        file_bytes: bytes,
        base_metadata: Dict[str, Any],
    ) -> List[Document]:
        """Extract text from DOCX and return semantic Document chunks."""
        try:
            import docx as python_docx
        except ImportError:
            logger.error("python-docx not installed — cannot chunk DOCX")
            return []

        try:
            doc = python_docx.Document(io.BytesIO(file_bytes))
        except Exception as e:
            logger.error(f"Failed to open DOCX: {e}")
            return []

        heading_stack: List[str] = []
        buffer: List[str] = []
        chunks: List[Document] = []

        def flush(force: bool = False) -> None:
            text = " ".join(buffer).strip()
            if not text:
                buffer.clear()
                return
            if not force and _approx_tokens(text) < MIN_TOKENS:
                return
            heading_path = " > ".join(heading_stack) if heading_stack else ""
            full_text = f"{heading_path}\n\n{text}".strip() if heading_path else text
            idx = len(chunks)
            meta = {
                **base_metadata,
                "chunk_index": idx,
                "heading_path": heading_path,
                "source": f"course_{base_metadata.get('course_id')}_module_{base_metadata.get('module_id')}_chunk_{idx}",
            }
            chunks.append(Document(page_content=full_text, metadata=meta))
            buffer.clear()

        # Process paragraphs
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style_name.startswith("Heading"):
                flush(force=True)
                # Extract level from style name e.g. "Heading 1" → 1
                m = re.search(r"(\d+)", style_name)
                level = int(m.group(1)) if m else 1
                heading_stack[:] = heading_stack[:level - 1]
                heading_stack.append(text)
            else:
                buffer.append(text)
                combined = " ".join(buffer)
                if _approx_tokens(combined) >= TARGET_TOKENS:
                    flush(force=True)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    buffer.append(f"Row: {row_text}")
                    combined = " ".join(buffer)
                    if _approx_tokens(combined) >= TARGET_TOKENS:
                        flush(force=True)

        flush(force=True)
        return chunks


# ─────────────────────────────────────────────────────────────────
# CourseRAGService
# ─────────────────────────────────────────────────────────────────

class CourseRAGService:
    """Manages per-course ChromaDB collections for Moodle course content.

    Each course gets its own Chroma collection named ``course_{course_id}``.
    The shared embeddings model (same as RAGService) is injected at construction
    time to avoid loading the model twice.
    """

    def __init__(
        self,
        embeddings: OpenAIEmbeddings,
        persist_directory: str,
        config_manager: Optional[ConfigurationManager] = None,
    ) -> None:
        self.embeddings = embeddings
        self.persist_directory = persist_directory
        self.config_manager = config_manager
        self.chunker = SemanticChunker()
        # Cache open Chroma collection handles keyed by course_id string
        self._collections: Dict[str, Chroma] = {}

        self._langid = translation_service.load_langid() if config_manager else None
        self._translation_llm = None
        if config_manager:
            try:
                self._translation_llm = translation_service.build_translation_llm(config_manager)
            except Exception as e:
                logger.error(f"Failed to initialize translation LLM: {e} — ingestion translation disabled")

        logger.info("CourseRAGService initialized")

    # ── Collection management ─────────────────────────────────────

    def _collection_name(self, course_id: str) -> str:
        return f"course_{course_id}"

    def _get_collection(self, course_id: str) -> Chroma:
        """Return (or lazily open) the Chroma collection for a course."""
        if course_id not in self._collections:
            self._collections[course_id] = Chroma(
                collection_name=self._collection_name(course_id),
                embedding_function=self.embeddings,
                persist_directory=self.persist_directory,
            )
            logger.info(f"Opened ChromaDB collection for course {course_id}")
        return self._collections[course_id]

    def delete_collection(self, course_id: str) -> None:
        """Drop the entire ChromaDB collection for a course."""
        name = self._collection_name(course_id)
        try:
            # Get the underlying chromadb client from the langchain wrapper
            store = self._get_collection(course_id)
            store._client.delete_collection(name)
            # Remove from cache
            self._collections.pop(course_id, None)
            logger.info(f"Deleted ChromaDB collection '{name}'")
        except Exception as e:
            logger.error(f"Failed to delete collection '{name}': {e}")
            raise

    # ── Ingestion ─────────────────────────────────────────────────

    def ingest_module(
        self,
        course_id: str,
        module_id: str,
        module_type: str,
        module_name: str,
        section_name: str,
        content_html: Optional[str] = None,
        content_raw_b64: Optional[str] = None,
        file_extension: Optional[str] = None,
    ) -> int:
        """Chunk and embed a course module into the per-course collection.

        Returns the number of chunks added.
        """
        base_metadata = {
            "type": "course_content",
            "course_id": course_id,
            "module_id": module_id,
            "module_type": module_type,
            "module_name": module_name,
            "section_name": section_name,
        }

        chunks: List[Document] = []

        if content_html:
            chunks = self.chunker.chunk_html(content_html, base_metadata)
        elif content_raw_b64 and file_extension:
            file_bytes = base64.b64decode(content_raw_b64)
            ext = file_extension.lower().lstrip(".")
            if ext == "pdf":
                chunks = self.chunker.chunk_pdf(file_bytes, base_metadata)
            elif ext in ("docx", "doc"):
                chunks = self.chunker.chunk_docx(file_bytes, base_metadata)
            else:
                logger.warning(f"Unsupported file extension '{ext}' — skipping")
                return 0
        else:
            logger.warning(f"No content provided for module {module_id} — skipping")
            return 0

        if not chunks:
            logger.info(f"No chunks produced for module {module_id}")
            return 0

        rag_config = self.config_manager.get_config().rag if self.config_manager else None
        chunks = self._translate_chunks_if_needed(chunks, rag_config)

        collection = self._get_collection(course_id)
        # Infomaniak embedding API accepts at most 99 items per call.
        batch_size = 99
        for i in range(0, len(chunks), batch_size):
            collection.add_documents(chunks[i : i + batch_size])
        logger.info(
            f"Indexed {len(chunks)} chunks for course {course_id} / module {module_id}"
        )
        return len(chunks)

    def _translate_chunks_if_needed(
        self, chunks: List[Document], rag_config: Optional[Any]
    ) -> List[Document]:
        """Translate every chunk of a non-French module to French.

        Language is detected once, from the module's first chunk — not per
        chunk — since a module is authored in one language. `page_content`
        (which already includes the heading breadcrumb baked in by
        SemanticChunker) is translated; `metadata["heading_path"]` is left
        untouched, since it's surfaced to the frontend as a citation/
        navigation breadcrumb back to the actual Moodle course structure,
        not used for retrieval.
        """
        if not chunks or rag_config is None or not rag_config.enable_ingestion_translation:
            return chunks
        if self._translation_llm is None:
            return chunks

        source_lang, should_translate = translation_service.decide_translation(
            chunks[0].page_content, self._langid,
            rag_config.langid_confidence_threshold, rag_config.min_langid_chars,
        )
        if not should_translate:
            return chunks

        out: List[Document] = []
        for chunk in chunks:
            prompt = translation_service.build_chunk_translation_prompt(chunk.page_content, source_lang)
            translated = translation_service.translate_to_french(prompt, self._translation_llm)
            new_meta = {**chunk.metadata, "source_language": source_lang}
            if translated:
                new_meta["original_text"] = chunk.page_content
                out.append(Document(page_content=translated, metadata=new_meta))
            else:
                out.append(Document(page_content=chunk.page_content, metadata=new_meta))
        return out

    def backfill_translations(self, course_id: str, rag_config: Any) -> Dict[str, int]:
        """Translate any chunk in this course's collection that predates the
        ingestion-translation feature (no `source_language` metadata key) to
        French, in place.

        Safe to re-run: chunks already tagged with `source_language`
        (translated, or confirmed French — a French chunk keeps no key, same
        convention as _translate_chunks_if_needed, so it's re-checked but
        never re-translated) are the only ones skipped outright, so an
        interrupted run can simply be invoked again. Uses ChromaDB's
        update_documents to replace page_content and re-embed in place,
        without touching chunk IDs or any other metadata.
        """
        stats = {"total": 0, "already_tagged": 0, "translated": 0, "unchanged_french": 0, "failed": 0}
        if self._translation_llm is None:
            return stats

        collection = self._get_collection(course_id)
        data = collection.get()
        ids = data.get("ids", [])
        documents = data.get("documents", [])
        metadatas = data.get("metadatas", [])
        stats["total"] = len(ids)

        update_ids: List[str] = []
        update_docs: List[Document] = []

        for doc_id, text, meta in zip(ids, documents, metadatas):
            if meta.get("source_language"):
                stats["already_tagged"] += 1
                continue

            source_lang, should_translate = translation_service.decide_translation(
                text or "", self._langid,
                rag_config.langid_confidence_threshold, rag_config.min_langid_chars,
            )
            if not should_translate:
                stats["unchanged_french"] += 1
                continue

            prompt = translation_service.build_chunk_translation_prompt(text, source_lang)
            translated = translation_service.translate_to_french(prompt, self._translation_llm)
            if not translated:
                stats["failed"] += 1
                continue

            new_meta = {**meta, "source_language": source_lang, "original_text": text}
            update_ids.append(doc_id)
            update_docs.append(Document(page_content=translated, metadata=new_meta))
            stats["translated"] += 1

        batch_size = 99
        for i in range(0, len(update_ids), batch_size):
            collection.update_documents(update_ids[i:i + batch_size], update_docs[i:i + batch_size])

        return stats

    def delete_module(self, course_id: str, module_id: str) -> int:
        """Remove all chunks belonging to a module from the course collection.

        Returns the number of chunks deleted.
        """
        try:
            collection = self._get_collection(course_id)
            results = collection.get(where={"module_id": module_id})
            ids = results.get("ids", [])
            if ids:
                collection.delete(ids=ids)
                logger.info(
                    f"Deleted {len(ids)} chunks for course {course_id} / module {module_id}"
                )
            return len(ids)
        except Exception as e:
            logger.error(f"Failed to delete module {module_id} from course {course_id}: {e}")
            raise

    # ── Retrieval ─────────────────────────────────────────────────

    def similarity_search(
        self, query: str, course_id: str, k: int = 5
    ) -> List[Document]:
        """MMR search within a single course's collection (embeds query internally)."""
        try:
            collection = self._get_collection(course_id)
            data = collection.get()
            if not data.get("ids"):
                logger.info(f"Course collection '{course_id}' is empty")
                return []

            results = collection.max_marginal_relevance_search(query, k=k)
            seen: set = set()
            unique: List[Document] = []
            for doc in results:
                src = doc.metadata.get("source", "")
                if src not in seen:
                    seen.add(src)
                    unique.append(doc)

            logger.info(f"Course search for '{course_id}' returned {len(unique)} results")
            return unique

        except Exception as e:
            logger.error(f"Course similarity search failed for course {course_id}: {e}")
            return []

    def _search_with_embedding(
        self, embedding: List[float], course_id: str, k: int = 5
    ) -> List[Document]:
        """Cosine-similarity search using a pre-computed query embedding.

        Avoids re-calling the embeddings API for each course by using the raw
        ChromaDB collection.query() which accepts ``query_embeddings`` directly.
        """
        try:
            collection = self._get_collection(course_id)
            raw = collection._collection.query(
                query_embeddings=[embedding],
                n_results=k,
                include=["documents", "metadatas"],
            )
            docs: List[Document] = []
            if raw.get("documents") and raw["documents"][0]:
                for text, meta in zip(raw["documents"][0], raw["metadatas"][0]):
                    docs.append(Document(page_content=text or "", metadata=meta or {}))
            return docs
        except Exception as e:
            logger.error(f"_search_with_embedding failed for course {course_id}: {e}")
            return []

    # ── Cross-course retrieval ────────────────────────────────────

    def _enumerate_populated_courses(self) -> List[str]:
        """Return course IDs of all non-empty ChromaDB course_* collections.

        Compatible with ChromaDB ≥0.6.0 where list_collections() returns
        List[str] (collection names) rather than List[Collection] objects.
        """
        try:
            # Access the underlying chromadb client via any open collection,
            # or open a temporary dummy collection to get the client reference.
            if self._collections:
                client = next(iter(self._collections.values()))._client
            else:
                dummy = self._get_collection("_probe")
                client = dummy._client

            raw = client.list_collections()
            # Chroma ≥0.6: returns List[str]; older versions: List[Collection]
            if raw and not isinstance(raw[0], str):
                names = [c.name for c in raw]
            else:
                names = list(raw)

            ids: List[str] = []
            for name in names:
                if not name.startswith("course_") or name == "course__probe":
                    continue
                cid = name[len("course_"):]
                col = self._get_collection(cid)
                if col.get().get("ids"):   # only include non-empty collections
                    ids.append(cid)

            # Clean up the probe collection if we created it
            if "_probe" in self._collections:
                try:
                    client.delete_collection("course__probe")
                except Exception:
                    pass
                self._collections.pop("_probe", None)

            return ids
        except Exception as e:
            logger.error(f"_enumerate_populated_courses failed: {e}")
            return []

    def similarity_search_all_courses(
        self,
        query: str,
        k_per_course: int = 1,
        priority_course_id: Optional[str] = None,
        allowed_course_ids: Optional[list] = None,
    ) -> List[Document]:
        """Query course collections the user is enrolled in.

        Embeds the query once, then searches every course collection using the
        pre-computed vector via the raw ChromaDB API.  This replaces N sequential
        embedding API calls (one per course) with a single call, cutting latency
        from O(N × embed_time) to O(embed_time + N × vector_search_time).

        The priority course gets k=6 results; all others get k_per_course results.

        If ``allowed_course_ids`` is provided, only those collections are queried.
        """
        all_docs: List[Document] = []
        course_ids = self._enumerate_populated_courses()
        if not course_ids:
            logger.info("similarity_search_all_courses: no populated courses found")
            return all_docs

        # Apply enrolment filter
        if allowed_course_ids is not None:
            allowed_set = set(str(cid) for cid in allowed_course_ids)
            course_ids = [cid for cid in course_ids if cid in allowed_set]
            if not course_ids:
                logger.info("similarity_search_all_courses: user enrolled in no indexed courses")
                return all_docs

        # Embed once and reuse across all collections
        try:
            embedding = self.embeddings.embed_query(query)
        except Exception as e:
            logger.error(f"similarity_search_all_courses: embedding failed: {e}")
            return all_docs

        # Always collect priority course results first so they are not pushed
        # out of the top-N cap by non-priority courses processed earlier.
        priority_docs: List[Document] = []
        other_docs: List[Document] = []

        for cid in course_ids:
            if cid == priority_course_id:
                docs = self._search_with_embedding(embedding, cid, k=6)
                priority_docs.extend(docs)
            else:
                docs = self._search_with_embedding(embedding, cid, k=k_per_course)
                other_docs.extend(docs)

        all_docs = priority_docs + other_docs
        logger.info(
            f"similarity_search_all_courses: {len(all_docs)} docs across "
            f"{len(course_ids)} courses (priority={priority_course_id})"
        )
        return all_docs

    # ── Status ────────────────────────────────────────────────────

    def get_course_status(self, course_id: str) -> Dict[str, Any]:
        """Return chunk and module counts for a course collection."""
        try:
            collection = self._get_collection(course_id)
            data = collection.get()
            ids = data.get("ids", [])
            metadatas = data.get("metadatas", [])
            module_ids = {m.get("module_id") for m in metadatas if m.get("module_id")}
            return {
                "course_id": course_id,
                "collection": self._collection_name(course_id),
                "chunk_count": len(ids),
                "module_count": len(module_ids),
            }
        except Exception as e:
            logger.error(f"Failed to get status for course {course_id}: {e}")
            return {
                "course_id": course_id,
                "collection": self._collection_name(course_id),
                "chunk_count": 0,
                "module_count": 0,
            }
